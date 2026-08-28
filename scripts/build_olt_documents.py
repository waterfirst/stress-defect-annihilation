from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "olt"
SOURCE = OUT / "manuscript_source.md"
FIGURE_DIR = OUT / "figures"


EQUATIONS = {
    "EQ1": "R = C(σ₁ − σ₂)d",
    "EQ2": "q₁ = σₓₓ − σᵧᵧ",
    "EQ3": "q₂ = 2σₓᵧ",
    "EQ4": "σ₁ − σ₂ = √(q₁² + q₂²)",
    "EQ5": "θ = ½ atan2(q₂, q₁)",
    "EQ6": "∂q/∂t = ∇·[Dₛ,max m(x,y)∇q] − [m(x,y)/τM]q",
    "EQ7": "Fo = Dₛ,max t/p²",
    "EQ8": "Da = p²/(Dₛ,max τM)",
    "EQ9": "∂q*/∂Fo = ∇*·(m∇*q*) − Da m q*",
    "EQ10": "t = Fo p²/Dₛ,max",
    "EQ11": "Cq(Fo) = ∫Ωs q² dA / ∫Ωs q₀² dA",
    "EQ12": "A(Fo) = A₀ exp[−(λ + Da)Fo]",
}


FIGURES = {
    "FIGURE_1": (
        FIGURE_DIR / "Figure_1_experimental_conditions.png",
        6.65,
        "Fig. 1. Experimental matrix and representative drilled-hole morphology. "
        "The retained source panel lists the 18 combinations of categorical pulse energy, pitch, and nominal target depth. "
        "Top-view and cross-sectional annotations are reproduced from the original project export without data-pixel alteration.",
    ),
    "FIGURE_2": (
        FIGURE_DIR / "Figure_2_energy_maps.png",
        5.85,
        "Fig. 2. Representative 550-nm total-retardance maps for Low, Middle, and High categorical pulse-energy settings at 20-µm pitch and nominal through-hole depth. "
        "The common displayed scale spans 0–200 nm. The panels support an ordinal comparison; absolute energy increments were not retained.",
    ),
    "FIGURE_3": (
        FIGURE_DIR / "Figure_3_depth_profiles.png",
        6.65,
        "Fig. 3. High-energy retardance maps and retained line profiles for nominal target depths of 500, 400, and 300 µm. "
        "The upper row uses 20-µm pitch (indices 3, 9, and 15), and the lower row uses 10-µm pitch (indices 6, 12, and 18). "
        "For both pitches, the 400-µm representative profile is the largest of the three tested depths.",
    ),
    "FIGURE_4": (
        FIGURE_DIR / "Figure_4_slow_axis_maps.png",
        6.65,
        "Fig. 4. Retained slow-axis maps from six laser-processing conditions. Line segments show the instrument-reported effective slow-axis orientation over total-retardance backgrounds. "
        "Individual colour limits differ among panels, so the figure is used to compare orientation texture rather than retardance magnitude.",
    ),
    "FIGURE_5": (
        FIGURE_DIR / "Figure_5_crack_correspondence.png",
        6.55,
        "Fig. 5. Qualitative crack-field correspondence from the retained project record: slow-axis map near adjacent holes, optical micrograph of a curved inter-hole crack, and a higher-resolution orientation map. "
        "This single example motivates registration-based validation but does not establish predictive accuracy.",
    ),
    "FIGURE_6": (
        FIGURE_DIR / "Figure_6_reduced_order_fields.png",
        6.55,
        "Fig. 6. Equal-scale reduced-order fields for opposite- and same-polarity controls at three Fourier exposures (h/p = 0.55, Da = 0.5). "
        "The early threshold favours the same-polarity control, whereas linear overlap leaves a smaller opposite-polarity tail at long exposure. "
        "The panels are deterministic Plotly outputs, not generated or altered research images.",
    ),
    "FIGURE_7": (
        FIGURE_DIR / "Figure_7_design_map.png",
        6.45,
        "Fig. 7. Dimensionless polarity and mobility-radius design map. "
        "The heat map shows the signed difference in quadratic field content between controls, and the neighbouring threshold plot shows that Fo90 is governed mainly by h/p rather than scalar polarity.",
    ),
    "FIGURE_8": (
        FIGURE_DIR / "Figure_8_verification_scaling.png",
        6.45,
        "Fig. 8. Numerical verification and scaling. "
        "Grid refinement reduces cosine-mode error, the selected Fo90 metric varies modestly over the tested Damköhler range, and conversion to seconds spans two decades when the uncalibrated effective mobility is varied by two decades.",
    ),
}


TABLE_1 = [
    ["Index", "Energy", "Pitch (µm)", "Target depth (µm)", "Nominal fraction", "State"],
    ["1", "Low", "20", "500", "100%", "Through-hole"],
    ["2", "Middle", "20", "500", "100%", "Through-hole"],
    ["3", "High", "20", "500", "100%", "Through-hole"],
    ["4", "Low", "10", "500", "100%", "Through-hole"],
    ["5", "Middle", "10", "500", "100%", "Through-hole"],
    ["6", "High", "10", "500", "100%", "Through-hole"],
    ["7", "Low", "20", "400", "80%", "Partial depth"],
    ["8", "Middle", "20", "400", "80%", "Partial depth"],
    ["9", "High", "20", "400", "80%", "Partial depth"],
    ["10", "Low", "10", "400", "80%", "Partial depth"],
    ["11", "Middle", "10", "400", "80%", "Partial depth"],
    ["12", "High", "10", "400", "80%", "Partial depth"],
    ["13", "Low", "20", "300", "60%", "Partial depth"],
    ["14", "Middle", "20", "300", "60%", "Partial depth"],
    ["15", "High", "20", "300", "60%", "Partial depth"],
    ["16", "Low", "10", "300", "60%", "Partial depth"],
    ["17", "Middle", "10", "300", "60%", "Partial depth"],
    ["18", "High", "10", "300", "60%", "Partial depth"],
]


TABLE_2 = [
    ["Quantity", "Definition/value", "Status"],
    ["p", "Hole-centre length scale", "Experimental scale; normalized in model"],
    ["Rv/p", "0.16", "Assumed geometry"],
    ["a/p", "0.08", "Assumed basis-field width"],
    ["h/p", "0.2–1.0; reference 0.55", "Swept mobility radius"],
    ["Da", "0–5; reference 0.5", "Swept dimensionless rate"],
    ["Domain", "6p × 6p", "Assumed zero-flux exterior"],
    ["Grid", "129 × 129 reference; 97 × 97 maps", "Verified against cosine mode"],
    ["ΔFo", "0.18(Δx/p)²", "Explicit stability setting"],
    ["Ds,max", "Not calibrated", "Required for time conversion"],
    ["τM", "Not calibrated", "Required for physical relaxation rate"],
]


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "B8B8B8", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    definitions = {
        "OLT Title": ("Arial", 16.5, True, "111111"),
        "OLT Author": ("Times New Roman", 11.2, False, "111111"),
        "OLT Affiliation": ("Times New Roman", 9.5, False, "444444"),
        "OLT Heading 1": ("Arial", 12.2, True, "222222"),
        "OLT Heading 2": ("Arial", 10.7, True, "222222"),
        "OLT Abstract": ("Times New Roman", 9.8, False, "111111"),
        "OLT Caption": ("Times New Roman", 8.8, False, "222222"),
        "OLT Equation": ("Cambria Math", 11.0, False, "111111"),
        "OLT Reference": ("Times New Roman", 8.6, False, "222222"),
    }
    for name, (font_name, size, bold, color) in definitions.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)

    styles["OLT Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["OLT Title"].paragraph_format.space_after = Pt(10)
    styles["OLT Author"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["OLT Author"].paragraph_format.space_after = Pt(2)
    styles["OLT Affiliation"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["OLT Affiliation"].paragraph_format.space_after = Pt(2)
    styles["OLT Heading 1"].paragraph_format.space_before = Pt(10)
    styles["OLT Heading 1"].paragraph_format.space_after = Pt(4)
    styles["OLT Heading 1"].paragraph_format.keep_with_next = True
    styles["OLT Heading 2"].paragraph_format.space_before = Pt(7)
    styles["OLT Heading 2"].paragraph_format.space_after = Pt(3)
    styles["OLT Heading 2"].paragraph_format.keep_with_next = True
    styles["OLT Abstract"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["OLT Abstract"].paragraph_format.line_spacing = 1.05
    styles["OLT Abstract"].paragraph_format.space_after = Pt(5)
    styles["OLT Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["OLT Caption"].paragraph_format.line_spacing = 1.0
    styles["OLT Caption"].paragraph_format.space_after = Pt(7)
    styles["OLT Equation"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["OLT Equation"].paragraph_format.space_before = Pt(3)
    styles["OLT Equation"].paragraph_format.space_after = Pt(3)
    styles["OLT Reference"].paragraph_format.left_indent = Inches(0.18)
    styles["OLT Reference"].paragraph_format.first_line_indent = Inches(-0.18)
    styles["OLT Reference"].paragraph_format.line_spacing = 1.0
    styles["OLT Reference"].paragraph_format.space_after = Pt(2)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(8)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_equation(doc: Document, key: str, number: int) -> None:
    p = doc.add_paragraph(style="OLT Equation")
    usable = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    p.paragraph_format.tab_stops.add_tab_stop(usable - Inches(0.05), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(EQUATIONS[key])
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    p.add_run(f"\t({number})")


def add_figure(doc: Document, key: str) -> None:
    path, width, caption = FIGURES[key]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="OLT Caption")
    cap.paragraph_format.keep_together = True
    add_inline(cap, caption)


def add_table(doc: Document, rows: list[list[str]], caption: str, widths: list[float]) -> None:
    cap = doc.add_paragraph(style="OLT Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    cap.add_run(caption).bold = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for i, row_values in enumerate(rows):
        row = table.rows[i]
        prevent_row_split(row)
        for j, value in enumerate(row_values):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if i == 0:
                set_cell_shading(cell, "E6E6E6")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or j != len(row_values) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(7.3 if len(rows) > 15 else 8.0)
            run.bold = i == 0
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_front_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "777777")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_manuscript() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.30)
    add_page_number(section)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    front_index = 0
    equation_number = 0
    current_heading = ""
    abstract_pending = False
    rule_added = False

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        token = line[1:-1] if line.startswith("[") and line.endswith("]") else line
        if line.startswith("# "):
            p = doc.add_paragraph(style="OLT Title")
            p.add_run(line[2:])
            front_index = 1
            continue
        if front_index and front_index < 5 and not line.startswith("#"):
            style = "OLT Author" if front_index == 1 else "OLT Affiliation"
            p = doc.add_paragraph(style=style)
            add_inline(p, line)
            front_index += 1
            if front_index == 4 and not rule_added:
                add_front_rule(doc)
                rule_added = True
            continue
        if line.startswith("## "):
            current_heading = line[3:]
            p = doc.add_paragraph(style="OLT Heading 1")
            p.add_run(current_heading)
            abstract_pending = current_heading == "Abstract"
            continue
        if line.startswith("### "):
            current_heading = line[4:]
            p = doc.add_paragraph(style="OLT Heading 2")
            p.add_run(current_heading)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.add_run("• ")
            add_inline(p, line[2:])
            continue
        if token in EQUATIONS:
            equation_number += 1
            add_equation(doc, token, equation_number)
            continue
        if token in FIGURES:
            add_figure(doc, token)
            continue
        if line == "[TABLE_1]":
            add_table(
                doc,
                TABLE_1,
                "Table 1. Factorial drilling matrix reconstructed from the original project record.",
                [0.50, 0.78, 0.86, 1.25, 1.05, 1.35],
            )
            continue
        if line == "[TABLE_2]":
            add_table(
                doc,
                TABLE_2,
                "Table 2. Reduced-order model quantities and calibration status.",
                [1.05, 2.35, 2.75],
            )
            continue
        if line.startswith("{REF}"):
            p = doc.add_paragraph(style="OLT Reference")
            add_inline(p, line[5:].strip())
            continue

        style = "OLT Abstract" if abstract_pending else "Normal"
        p = doc.add_paragraph(style=style)
        add_inline(p, line)
        if abstract_pending:
            abstract_pending = False
        if line.startswith("**Keywords:"):
            p.paragraph_format.space_after = Pt(6)
        if current_heading in {
            "CRediT authorship contribution statement",
            "Funding",
            "Declaration of competing interest",
            "Data availability",
            "Declaration of generative AI and AI-assisted technologies in the writing process",
        }:
            p.paragraph_format.keep_together = True

    path = OUT / "manuscript_OLT.docx"
    doc.core_properties.title = "Depth-dependent photoelastic retardance around picosecond Bessel-beam-drilled glass and a calibration-ready reduced-order model"
    doc.core_properties.subject = "Optics & Laser Technology full-length research article"
    doc.core_properties.author = "Nakcho Choi; Jeongjin Park"
    doc.core_properties.keywords = "Bessel beam, glass drilling, photoelasticity, retardance, residual stress"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_manuscript())
