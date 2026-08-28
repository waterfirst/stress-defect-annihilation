from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "olt"
SOURCE = OUT / "manuscript_source.md"
DOCX = OUT / "manuscript_OLT.docx"


def section_text(text: str, start: str, end: str) -> str:
    return text.split(start, 1)[1].split(end, 1)[0].strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\wµ–−]+\b", text, flags=re.UNICODE))


def expand_citations(body: str) -> set[int]:
    result: set[int] = set()
    for block in re.findall(r"\[([0-9,–\- ]+)\]", body):
        for item in block.split(","):
            item = item.strip()
            if not item:
                continue
            parts = re.split(r"[–\-]", item)
            if len(parts) == 2:
                start, end = map(int, parts)
                result.update(range(start, end + 1))
            else:
                result.add(int(parts[0]))
    return result


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    abstract = section_text(text, "## Abstract", "**Keywords:")
    keywords_line = re.search(r"\*\*Keywords:\*\*\s*(.+)", text)
    if keywords_line is None:
        raise AssertionError("Keywords missing")
    keywords = [item.strip() for item in keywords_line.group(1).split(";")]
    highlights = [
        line[2:].strip()
        for line in section_text(text, "## Highlights", "## 1. Introduction").splitlines()
        if line.startswith("- ")
    ]
    body = text.split("## References", 1)[0]
    references = [line for line in text.splitlines() if line.startswith("{REF}")]
    cited = expand_citations(body)
    expected = set(range(1, len(references) + 1))

    assert word_count(abstract) <= 250, word_count(abstract)
    assert 1 <= len(keywords) <= 7, len(keywords)
    assert 3 <= len(highlights) <= 5, len(highlights)
    assert all(len(item) <= 85 for item in highlights), [(len(item), item) for item in highlights]
    assert cited == expected, {"missing": sorted(expected - cited), "extra": sorted(cited - expected)}
    assert "Seongjun Lee" not in text
    assert DOCX.exists(), DOCX

    doc = Document(DOCX)
    assert len(doc.inline_shapes) == 8, len(doc.inline_shapes)
    assert len(doc.tables) == 2, len(doc.tables)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Seongjun Lee" not in full_text
    assert doc.core_properties.author == "Nakcho Choi; Jeongjin Park", doc.core_properties.author
    assert "Liquid crystals disclination images" not in full_text
    assert "predicts crack path" not in full_text.lower()
    assert "maximum at 75%" not in full_text.lower()
    assert "MPa" not in abstract

    print(f"abstract_words={word_count(abstract)}")
    print(f"keywords={len(keywords)}")
    print("highlight_chars=" + ",".join(str(len(item)) for item in highlights))
    print(f"references={len(references)} cited={len(cited)}")
    print(f"docx_figures={len(doc.inline_shapes)} tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
